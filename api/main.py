import os
import json
import time
from pathlib import Path
import socket
from typing import Any, Dict, List, Optional, Tuple
import hashlib
import platform
import json as _json
import asyncio
from datetime import datetime
import logging

import numpy as np
import requests
import aiohttp
from fastapi import Depends, FastAPI, Header, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

try:
    # OpenAI SDK v1
    from openai import OpenAI  # type: ignore
except Exception:  # pragma: no cover
    OpenAI = None  # type: ignore

# Import PostgreSQL memory system
try:
    from src.layers.memory_system import MemorySystem
    POSTGRES_MEMORY_AVAILABLE = True
except ImportError:
    POSTGRES_MEMORY_AVAILABLE = False
    MemorySystem = None

API_KEY_ENV = os.getenv("API_KEY", "demo_key_123")

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def require_api_key(x_api_key: Optional[str] = Header(None)) -> None:
    if API_KEY_ENV and x_api_key != API_KEY_ENV:
        raise HTTPException(status_code=401, detail="Invalid API key")


app = FastAPI(title="Capstone Demo API", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)


# Identity headers middleware
@app.middleware("http")
async def add_identity_headers(request, call_next):
    response = await call_next(request)
    try:
        response.headers["X-Api-Hostname"] = HOSTNAME
        response.headers["X-Api-Started-At"] = str(STARTED_AT)
        response.headers["X-Api-Image-Tag"] = IMAGE_TAG
        response.headers["X-Api-Commit"] = COMMIT_SHA
        response.headers["X-Api-Instance-Id"] = INSTANCE_ID
    except Exception:
        pass
    return response


# ---------------------------
# Basic health and metrics
# ---------------------------
def _load_identity_from_file() -> Dict[str, Any]:
    """Optionally load identity from a JSON file specified by IDENTITY_FILE or ./identity.json.
    Values in this file augment environment-derived identity.
    """
    locations = []
    env_path = os.getenv("IDENTITY_FILE")
    if env_path:
        locations.append(Path(env_path))
    # next to this module
    locations.append(Path(__file__).resolve().parent / "identity.json")
    # project root (two levels up from api/main.py)
    locations.append(Path(__file__).resolve().parents[1] / "identity.json")
    for p in locations:
        try:
            if p.exists():
                return _json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
    return {}


_IDENTITY_FILE_DATA = _load_identity_from_file()


def build_identity() -> Dict[str, Any]:
    """Build a standardized identity object for the API instance.

    Priority order per field: file override → environment → computed defaults.
    """
    service = _IDENTITY_FILE_DATA.get("service") or os.getenv("SERVICE_NAME", "capstone-demo-api")
    version = _IDENTITY_FILE_DATA.get("version") or os.getenv("SERVICE_VERSION", "0.1.0")
    # Select a human-visible identificator if provided, else use instance hash
    identificator = (
        _IDENTITY_FILE_DATA.get("identificator")
        or os.getenv("API_IDENTIFICATOR")
        or INSTANCE_ID
    )
    identity: Dict[str, Any] = {
        "service": service,
        "version": version,
        "hostname": HOSTNAME,
        "started_at": STARTED_AT,
        "image_tag": _IDENTITY_FILE_DATA.get("image_tag") or IMAGE_TAG,
        "commit_sha": _IDENTITY_FILE_DATA.get("commit_sha") or COMMIT_SHA,
        "build_time": _IDENTITY_FILE_DATA.get("build_time") or BUILD_TIME,
        "instance_id": INSTANCE_ID,
        "identificator": identificator,
        "python": platform.python_version(),
        "pid": os.getpid(),
    }
    # Allow custom fields from file under "extras"
    extras = _IDENTITY_FILE_DATA.get("extras")
    if isinstance(extras, dict):
        identity["extras"] = extras
    return identity


@app.get("/identity")
def identity(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return build_identity()


@app.get("/health")
def health(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"status": "ok", "instance": build_identity()}


@app.get("/performance/metrics")
def performance_metrics(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {
        "overall": {"success_rate": 99.5, "avg_response_time": 12.3, "total_requests": 1234},
        "tools": {
            "ping": {"success_rate": 100.0, "avg_response_time": 5.1, "total_requests": 100},
            "list_files": {"success_rate": 99.0, "avg_response_time": 9.9, "total_requests": 50},
        },
    }


@app.get("/performance/alerts")
def performance_alerts(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"alerts": []}


@app.get("/performance/dashboard")
def performance_dashboard(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"dashboard": {"active_tools": 5, "uptime_hours": 12.5, "error_rate": 0.05}}


@app.get("/performance/health")
def performance_health(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"status": "ok", "metrics": {"cpu_usage": 12.1, "memory_usage": 43.2}}


# Tools registry
@app.get("/tools/list")
def tools_list(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    # Reverted to canonical tool names (including get_system_status)
    return {"tools": ["ping", "list_files", "read_file", "analyze_code", "get_system_status"]}


@app.get("/tools/registry/health")
def tools_registry_health(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"status": "ok", "registered": 5}


class ToolCall(BaseModel):
    name: str
    arguments: Dict[str, Any] = {}


@app.post("/api/v1/tools/call")
def tools_call(body: ToolCall, _: None = Depends(require_api_key)) -> Dict[str, Any]:
    if body.name == "ping":
        # Include lightweight components so UI can show instance + LLM + Memory + Endpoints
        components = _components_snapshot()
        return {
            "content": [{"type": "text", "text": "pong"}],
            "instance": build_identity(),
            "llm": components.get("llm", {}),
            "memory": components.get("memory", {}),
            "endpoints": components.get("endpoints", {}),
        }
    if body.name == "list_files":
        try:
            limit = int(body.arguments.get("limit", 20))
        except Exception:
            limit = 20
        # Reuse internal helper
        import asyncio
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            data = loop.run_until_complete(_list_documents(limit=limit))
        finally:
            loop.close()
        return {"tool": "list_files", "data": data}
    if body.name == "read_file":
        # Expect argument: id
        doc_id = body.arguments.get("id")
        import asyncio
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            data: Dict[str, Any]
            data = {"success": False, "error": "invalid_id"}
            if doc_id is not None:
                try:
                    ms = loop.run_until_complete(_get_simple_memory_system())
                    if ms:
                        doc = loop.run_until_complete(ms.get_document_by_id(int(doc_id)))
                        if doc:
                            data = {"success": True, "document": doc, "type": "postgresql_pgvector"}
                        else:
                            data = {"success": False, "error": "not_found"}
                    else:
                        # Fallback to JSON by index
                        mem = _load_memory()
                        items = mem.get("items", [])
                        idx = max(0, min(len(items) - 1, int(doc_id) - 1))
                        if 0 <= idx < len(items):
                            it = items[idx]
                            data = {"success": True, "document": {"id": str(doc_id), "source": "json_file", "content": it.get("response", "")}, "type": "json_file"}
                        else:
                            data = {"success": False, "error": "not_found"}
                except Exception as e:
                    data = {"success": False, "error": str(e)}
        finally:
            loop.close()
        return {"tool": "read_file", "data": data}
    if body.name == "analyze_code":
        # Use preview/analyze logic (mock code analyzer)
        try:
            file_path = str(body.arguments.get("file_path", ""))
            res = preview_analyze(file_path, None)  # type: ignore[arg-type]
            return {"tool": "analyze_code", "data": res}
        except Exception as e:
            return {"tool": "analyze_code", "error": str(e)}
    if body.name == "get_system_status":
        # Aggregate basic status for UI diagnostics
        try:
            return {
                "tool": "get_system_status",
                "instance": build_identity(),
                "components": _components_snapshot(),
            }
        except Exception as e:
            return {"tool": "get_system_status", "error": str(e)}
    return {"content": [{"type": "text", "text": f"Executed {body.name} with {body.arguments}"}]}


# ---------------------------
# Simple RAG + LLM wiring
# ---------------------------

# Memory system configuration
MEMORY_PATH = Path(os.getenv("MEMORY_PATH", "./memory_store.json")).resolve()
EMBED_MODEL = os.getenv("EMBED_MODEL", "text-embedding-3-small")
CHAT_MODEL = os.getenv("CHAT_MODEL", "gpt-4o-mini")

# Provider switch: "ollama" or default "openai"
LLM_PROVIDER = os.getenv("LLM_PROVIDER", "openai").strip().lower()
OLLAMA_BASE = os.getenv("OLLAMA_BASE", "http://lfc-ollama:11434").rstrip("/")
OLLAMA_CHAT_MODEL = os.getenv("OLLAMA_CHAT_MODEL", "llama3.2:3b")
OLLAMA_EMBED_MODEL = os.getenv("OLLAMA_EMBED_MODEL", "mxbai-embed-large")
EMBED_DIM = int(os.getenv("EMBED_DIM", "1024"))

# Simple PostgreSQL memory system
import asyncpg
from typing import List, Dict, Any
import re

# Import Interaction class
try:
    from src.layers.memory_system import Interaction
except ImportError:
    # Fallback if import fails
    from dataclasses import dataclass
    from datetime import datetime
    
    @dataclass
    class Interaction:
        user_input: str
        response: str
        tool_calls: List[Any]
        metadata: Dict[str, Any]
        timestamp: datetime
        source: str = "unknown"

class SimplePostgreSQLMemory:
    """Simple PostgreSQL memory system for direct database access"""
    
    def __init__(self, connection_string: str):
        self.connection_string = connection_string
        self._pool = None
    
    async def initialize(self):
        """Initialize connection pool"""
        try:
            self._pool = await asyncpg.create_pool(
                self.connection_string,
                min_size=1,
                max_size=5,
                command_timeout=15
            )
            
            # Create chat_history table if it doesn't exist
            await self._create_chat_history_table()
            
            # Create main schema including vector_store table
            await self._create_schema()
            
            return True
        except Exception as e:
            print(f"Failed to initialize PostgreSQL pool: {e}")
            return False
    
    async def _create_chat_history_table(self):
        """Create chat_history table if it doesn't exist"""
        try:
            async with self._pool.acquire() as conn:
                # Create chat history table
                await conn.execute("""
                    CREATE TABLE IF NOT EXISTS chat_history (
                        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                        session_id VARCHAR(255) NOT NULL,
                        user_message TEXT NOT NULL,
                        assistant_response TEXT NOT NULL,
                        context_docs JSONB DEFAULT '[]',
                        used_docs INTEGER DEFAULT 0,
                        metadata JSONB DEFAULT '{}',
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        response_time_ms INTEGER,
                        model_used VARCHAR(100),
                        streamed BOOLEAN DEFAULT FALSE
                    )
                """)
                
                # Create indexes
                await conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chat_history_session_id 
                    ON chat_history (session_id)
                """)
                
                await conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chat_history_created_at 
                    ON chat_history (created_at)
                """)
                
                await conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_chat_history_session_created 
                    ON chat_history (session_id, created_at)
                """)
                
                print("Chat history table and indexes created successfully")
                
        except Exception as e:
            print(f"Failed to create chat history table: {e}")
    
    async def _create_schema(self):
        """Create the main database schema including vector_store table"""
        try:
            async with self._pool.acquire() as conn:
                # Enable pgvector extension
                await conn.execute("CREATE EXTENSION IF NOT EXISTS vector")
                
                # Create vector_store table
                await conn.execute("""
                    CREATE TABLE IF NOT EXISTS vector_store (
                        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                        source VARCHAR(500) NOT NULL,
                        content TEXT NOT NULL,
                        embedding vector(1024),
                        metadata JSONB DEFAULT '{}',
                        created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
                        updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
                    )
                """)
                
                # Create indexes for vector_store
                await conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_vector_store_source 
                    ON vector_store (source)
                """)
                
                await conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_vector_store_created_at 
                    ON vector_store (created_at)
                """)
                
                await conn.execute("""
                    CREATE INDEX IF NOT EXISTS idx_vector_store_embedding 
                    ON vector_store USING ivfflat (embedding vector_cosine_ops)
                    WITH (lists = 100)
                """)
                
                print("Vector store table and indexes created successfully")
                
        except Exception as e:
            print(f"Failed to create vector store schema: {e}")
    
    async def query_similar(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """Simple text-based search (fallback when embeddings not available)"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Attempt search on content and source
                async def _run(q: str) -> List[asyncpg.Record]:
                    return await conn.fetch(
                        """
                        SELECT id, source, content, created_at 
                        FROM vector_store 
                        WHERE content ILIKE $1 OR source ILIKE $1
                        ORDER BY created_at DESC 
                        LIMIT $2
                        """,
                        f"%{q}%", k
                    )

                results = await _run(query)
                # Filename normalization fallback: .txt <-> .md
                if not results:
                    if re.search(r"\.txt$", query, flags=re.IGNORECASE):
                        alt = re.sub(r"\.txt$", ".md", query, flags=re.IGNORECASE)
                        results = await _run(alt)
                    elif re.search(r"\.md$", query, flags=re.IGNORECASE):
                        alt = re.sub(r"\.md$", ".txt", query, flags=re.IGNORECASE)
                        results = await _run(alt)

                return [
                    {
                        "id": str(r["id"]),
                        "source": r["source"],
                        "content": r["content"],
                        "created_at": r["created_at"].isoformat() if r["created_at"] else None
                    }
                    for r in results
                ]
        except Exception as e:
            print(f"PostgreSQL query failed: {e}")
            return []
    
    async def query_similar_with_embedding(self, query: str, embedding: List[float], k: int = 5) -> Dict[str, Any]:
        """Semantic search using embeddings (for chat system)"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Extract meaningful search terms from the query
                search_terms = self._extract_search_terms(query)
                
                # Try each search term
                results = []
                for term in search_terms:
                    if not term.strip():
                        continue
                    
                    term_results = await conn.fetch(
                        """
                        SELECT id, source, content, created_at 
                        FROM vector_store 
                        WHERE content ILIKE $1 OR source ILIKE $1
                        ORDER BY created_at DESC 
                        LIMIT $2
                        """,
                        f"%{term}%", k
                    )
                    
                    # Add unique results
                    for r in term_results:
                        if not any(existing["id"] == str(r["id"]) for existing in results):
                            results.append(r)
                    
                    if len(results) >= k:
                        break
                
                # Filename normalization fallback for specific file queries
                if not results and any(term.lower().endswith(('.txt', '.md', '.py', '.js')) for term in search_terms):
                    for term in search_terms:
                        if term.lower().endswith('.txt'):
                            alt = term[:-4] + '.md'
                            alt_results = await conn.fetch(
                                """
                                SELECT id, source, content, created_at 
                                FROM vector_store 
                                WHERE content ILIKE $1 OR source ILIKE $1
                                ORDER BY created_at DESC 
                                LIMIT $2
                                """,
                                f"%{alt}%", k
                            )
                            results.extend(alt_results)
                        elif term.lower().endswith('.md'):
                            alt = term[:-3] + '.txt'
                            alt_results = await conn.fetch(
                                """
                                SELECT id, source, content, created_at 
                                FROM vector_store 
                                WHERE content ILIKE $1 OR source ILIKE $1
                                ORDER BY created_at DESC 
                                LIMIT $2
                                """,
                                f"%{alt}%", k
                            )
                            results.extend(alt_results)
                
                # Limit results and format
                results = results[:k]
                snippets = []
                for r in results:
                    content = r["content"] or ""
                    # Create a preview (first 200 chars)
                    preview = content[:200] + "..." if len(content) > 200 else content
                    snippets.append({
                        "id": str(r["id"]),
                        "source": r["source"],
                        "content": content,
                        "preview": preview,
                        "created_at": r["created_at"].isoformat() if r["created_at"] else None
                    })
                
                return {
                    "success": True,
                    "snippets": snippets,
                    "total_found": len(snippets)
                }
                
        except Exception as e:
            print(f"PostgreSQL semantic query failed: {e}")
            return {"success": False, "snippets": [], "total_found": 0}
    
    def _extract_search_terms(self, query: str) -> List[str]:
        """Extract meaningful search terms from a user query"""
        # Remove common question words and extract key terms
        question_words = {
            'what', 'where', 'when', 'who', 'why', 'how', 'tell', 'me', 'about', 'the', 'a', 'an',
            'is', 'are', 'was', 'were', 'do', 'does', 'did', 'can', 'could', 'would', 'should',
            'file', 'files', 'document', 'documents', 'show', 'list', 'find', 'search'
        }
        
        # Split query into words and filter
        words = query.lower().split()
        terms = [word for word in words if word not in question_words and len(word) > 2]
        
        # Add original query as fallback
        if terms:
            terms.append(query)
        else:
            terms = [query]
        
        return terms
    
    async def get_document_count(self) -> int:
        """Get total document count"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                return await conn.fetchval("SELECT COUNT(*) FROM vector_store")
        except Exception as e:
            print(f"Failed to get document count: {e}")
            return 0
    
    async def get_sample_documents(self, limit: int = 5) -> List[Dict[str, Any]]:
        """Get sample documents for preview"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                results = await conn.fetch(
                    """
                    SELECT id, source, content, created_at 
                    FROM vector_store 
                    ORDER BY created_at DESC 
                    LIMIT $1
                    """,
                    limit
                )
                
                return [
                    {
                        "id": str(r["id"]),
                        "source": r["source"],
                        "content": r["content"][:200] + "..." if len(r["content"]) > 200 else r["content"],
                        "created_at": r["created_at"].isoformat() if r["created_at"] else None
                    }
                    for r in results
                ]
        except Exception as e:
            print(f"Failed to get sample documents: {e}")
            return []

    async def get_document_by_id(self, doc_id: int) -> Optional[Dict[str, Any]]:
        """Fetch full document by id."""
        try:
            if not self._pool:
                await self.initialize()
            async with self._pool.acquire() as conn:
                row = await conn.fetchrow(
                    """
                    SELECT id, source, content, created_at
                    FROM vector_store
                    WHERE id = $1
                    """,
                    doc_id,
                )
                if not row:
                    return None
                return {
                    "id": str(row["id"]),
                    "source": row["source"],
                    "content": row["content"],
                    "created_at": row["created_at"].isoformat() if row["created_at"] else None,
                }
        except Exception as e:
            print(f"Failed to get document by id: {e}")
            return None
    
    async def store_interaction(self, interaction: Interaction, embedding: List[float]) -> Dict[str, Any]:
        """Store interaction with embedding in PostgreSQL"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Store the interaction in vector_store table
                new_id = await conn.fetchval(
                    """
                    INSERT INTO vector_store (source, content, created_at)
                    VALUES ($1, $2, $3)
                    RETURNING id
                    """,
                    (
                        (getattr(interaction, 'metadata', {}).get('filename') if getattr(interaction, 'metadata', None) else None)
                        or getattr(interaction, 'source', 'unknown')
                    ),
                    f"{interaction.user_input}\n{interaction.response}",
                    datetime.now()
                )
                return {"success": True, "record_id": int(new_id)}
        except Exception as e:
            print(f"Failed to store interaction in PostgreSQL: {e}")
            return {"success": False, "error": str(e)}

    async def store_chat_history(self, session_id: str, user_message: str, 
                                assistant_response: str, context_docs: List[Dict[str, Any]] = None,
                                used_docs: int = 0, metadata: Dict[str, Any] = None,
                                response_time_ms: int = None, model_used: str = None,
                                streamed: bool = False) -> Dict[str, Any]:
        """Store chat conversation in the database"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Store chat history
                result = await conn.fetchrow("""
                    INSERT INTO chat_history 
                    (session_id, user_message, assistant_response, context_docs, used_docs, 
                     metadata, response_time_ms, model_used, streamed)
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9)
                    RETURNING id, created_at
                """, session_id, user_message, assistant_response, 
                     json.dumps(context_docs or []), used_docs,
                     json.dumps(metadata or {}), response_time_ms, model_used, streamed)
                
                return {
                    "success": True,
                    "chat_id": str(result['id']),
                    "created_at": result['created_at'].isoformat(),
                    "session_id": session_id
                }
                
        except Exception as e:
            print(f"Failed to store chat history: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    async def get_chat_history(self, session_id: str = None, limit: int = 50, 
                              offset: int = 0, include_context: bool = False) -> Dict[str, Any]:
        """Retrieve chat history from the database"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Build query based on parameters
                if session_id:
                    # Get chat history for specific session
                    query = """
                        SELECT id, session_id, user_message, assistant_response, 
                               context_docs, used_docs, metadata, created_at, 
                               response_time_ms, model_used, streamed
                        FROM chat_history 
                        WHERE session_id = $1
                        ORDER BY created_at DESC
                        LIMIT $2 OFFSET $3
                    """
                    params = [session_id, limit, offset]
                else:
                    # Get all chat history
                    query = """
                        SELECT id, session_id, user_message, assistant_response, 
                               context_docs, used_docs, metadata, created_at, 
                               response_time_ms, model_used, streamed
                        FROM chat_history 
                        ORDER BY created_at DESC
                        LIMIT $1 OFFSET $2
                    """
                    params = [limit, offset]
                
                # Execute query
                rows = await conn.fetch(query, *params)
                
                # Get total count for pagination
                if session_id:
                    total_count = await conn.fetchval(
                        "SELECT COUNT(*) FROM chat_history WHERE session_id = $1",
                        session_id
                    )
                else:
                    total_count = await conn.fetchval("SELECT COUNT(*) FROM chat_history")
                
                # Process results
                chats = []
                for row in rows:
                    chat = {
                        "id": str(row['id']),
                        "session_id": row['session_id'],
                        "user_message": row['user_message'],
                        "assistant_response": row['assistant_response'],
                        "used_docs": row['used_docs'],
                        "created_at": row['created_at'].isoformat(),
                        "response_time_ms": row['response_time_ms'],
                        "model_used": row['model_used'],
                        "streamed": row['streamed']
                    }
                    
                    if include_context:
                        chat["context_docs"] = row['context_docs']
                        chat["metadata"] = row['metadata']
                    
                    chats.append(chat)
                
                return {
                    "success": True,
                    "chats": chats,
                    "total_count": total_count,
                    "limit": limit,
                    "offset": offset,
                    "has_more": (offset + limit) < total_count
                }
                
        except Exception as e:
            print(f"Failed to retrieve chat history: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    async def get_session_summary(self, session_id: str) -> Dict[str, Any]:
        """Get a summary of a chat session"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Get session statistics
                stats = await conn.fetchrow("""
                    SELECT 
                        COUNT(*) as total_messages,
                        MIN(created_at) as first_message,
                        MAX(created_at) as last_message,
                        AVG(response_time_ms) as avg_response_time,
                        SUM(used_docs) as total_docs_used,
                        COUNT(CASE WHEN streamed = true THEN 1 END) as streamed_messages
                    FROM chat_history 
                    WHERE session_id = $1
                """, session_id)
                
                if not stats:
                    return {
                        "success": False,
                        "error": "Session not found"
                    }
                
                # Get recent messages for context
                recent_messages = await conn.fetch("""
                    SELECT user_message, assistant_response, created_at
                    FROM chat_history 
                    WHERE session_id = $1
                    ORDER BY created_at DESC
                    LIMIT 5
                """, session_id)
                
                return {
                    "success": True,
                    "session_id": session_id,
                    "statistics": {
                        "total_messages": stats['total_messages'],
                        "first_message": stats['first_message'].isoformat() if stats['first_message'] else None,
                        "last_message": stats['last_message'].isoformat() if stats['last_message'] else None,
                        "avg_response_time_ms": round(stats['avg_response_time'], 2) if stats['avg_response_time'] else None,
                        "total_docs_used": stats['total_docs_used'],
                        "streamed_messages": stats['streamed_messages']
                    },
                    "recent_messages": [
                        {
                            "user_message": msg['user_message'][:100] + "..." if len(msg['user_message']) > 100 else msg['user_message'],
                            "assistant_response": msg['assistant_response'][:100] + "..." if len(msg['assistant_response']) > 100 else msg['assistant_response'],
                            "created_at": msg['created_at'].isoformat()
                        }
                        for msg in recent_messages
                    ]
                }
                
        except Exception as e:
            print(f"Failed to get session summary: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    async def get_health_status(self) -> Dict[str, Any]:
        """Get health status of the memory system"""
        try:
            if not self._pool:
                await self.initialize()
            
            async with self._pool.acquire() as conn:
                # Check if tables exist
                tables_check = await conn.fetch("""
                    SELECT table_name 
                    FROM information_schema.tables 
                    WHERE table_schema = 'public' 
                    AND table_name IN ('vector_store', 'chat_history')
                """)
                
                existing_tables = [row['table_name'] for row in tables_check]
                
                # Check document counts
                vector_count = await conn.fetchval("SELECT COUNT(*) FROM vector_store")
                chat_count = await conn.fetchval("SELECT COUNT(*) FROM chat_history")
                
                # Check pgvector extension
                extension_check = await conn.fetchval("""
                    SELECT 1 FROM pg_extension WHERE extname = 'vector'
                """)
                
                return {
                    "status": "healthy" if self._pool and len(existing_tables) >= 2 else "unhealthy",
                    "connection_pool": "available" if self._pool else "unavailable",
                    "tables": {
                        "vector_store": "vector_store" in existing_tables,
                        "chat_history": "chat_history" in existing_tables
                    },
                    "pgvector_extension": bool(extension_check),
                    "document_counts": {
                        "vector_store": vector_count,
                        "chat_history": chat_count
                    },
                    "timestamp": datetime.now().isoformat()
                }
                
        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }

# Global memory system instance
_simple_memory_system: Optional[SimplePostgreSQLMemory] = None
_memory_system_lock = asyncio.Lock()

async def _get_simple_memory_system() -> Optional[SimplePostgreSQLMemory]:
    """Get or create simple PostgreSQL memory system"""
    global _simple_memory_system
    
    if _simple_memory_system is None:
        async with _memory_system_lock:
            if _simple_memory_system is None:
                try:
                    _simple_memory_system = SimplePostgreSQLMemory(
                        connection_string=os.getenv("DATABASE_URL")
                    )
                    await _simple_memory_system.initialize()
                except Exception as e:
                    print(f"Warning: Failed to initialize simple PostgreSQL memory system: {e}")
                    return None
    
    return _simple_memory_system

def _get_memory_system_sync() -> Optional[SimplePostgreSQLMemory]:
    """Synchronous wrapper for getting memory system"""
    try:
        # Check if we're in an async context
        import asyncio
        try:
            asyncio.get_running_loop()
            # We're in async context, return None to avoid conflicts
            return None
        except RuntimeError:
            # No event loop, safe to create one
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                result = loop.run_until_complete(_get_simple_memory_system())
                return result
            finally:
                loop.close()
    except Exception as e:
        print(f"Failed to get memory system synchronously: {e}")
        return None

# ---------------------------
# Tool helpers
# ---------------------------
async def _list_documents(limit: int = 20) -> Dict[str, Any]:
    """Return list of documents from PostgreSQL if available, else JSON fallback."""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            docs = await mem_system.get_sample_documents(limit=limit)
            total = await mem_system.get_document_count()
            return {
                "success": True,
                "type": "postgresql_pgvector",
                "total_documents": total,
                "documents": docs,
            }
    except Exception:
        # Fall back below
        pass

    # Fallback to JSON memory file
    mem = _load_memory()
    items = mem.get("items", [])
    documents: List[Dict[str, Any]] = []
    for idx, item in enumerate(items[:limit]):
        documents.append(
            {
                "id": str(idx + 1),
                "source": "json_file",
                "content": (item.get("response") or "")[:200],
                "created_at": None,
            }
        )
    return {
        "success": True,
        "type": "json_file",
        "total_documents": len(items),
        "documents": documents,
    }

async def _handle_chat_tools(message: str) -> Optional[Dict[str, Any]]:
    """Very lightweight tool router for chat messages.

    Triggers:
    - Explicit command: !list_docs [limit]
    - Heuristic: user asks to list/what files/documents present
    Returns a dict with keys: tool, text, data; or None if no tool is used.
    """
    msg = (message or "").lower()

    # Explicit command parsing
    if msg.startswith("!list_docs") or any(
        phrase in msg for phrase in [
            "what files are present",
            "what documents are present",
            "what files do we have",
            "what documents do we have",
            "list files",
            "list documents",
        ]
    ):
        # Optional limit parsing for explicit command: !list_docs 50
        parts = msg.split()
        limit = 20
        if len(parts) >= 2 and parts[0] == "!list_docs":
            try:
                limit = max(1, min(200, int(parts[1])))
            except Exception:
                limit = 20

        data = await _list_documents(limit=limit)
        docs = data.get("documents", [])
        lines: List[str] = []
        for d in docs:
            lines.append(
                f"- id {d.get('id')}, source {d.get('source')}, created_at {d.get('created_at')}"
            )
        text = (
            "Here are the documents currently in memory (top "
            f"{len(docs)}):\n" + ("\n".join(lines) if lines else "(none)")
        )
        return {"tool": "list_documents", "text": text, "data": data}

    return None

# Process identity
STARTED_AT = int(time.time())
HOSTNAME = os.getenv("HOSTNAME") or socket.gethostname()
IMAGE_TAG = os.getenv("IMAGE_TAG", "local")
COMMIT_SHA = os.getenv("COMMIT_SHA", "dev")
BUILD_TIME = os.getenv("BUILD_TIME", "")
INSTANCE_ID = hashlib.sha1(
    f"{HOSTNAME}|{STARTED_AT}|{IMAGE_TAG}|{COMMIT_SHA}|{BUILD_TIME}".encode("utf-8", errors="ignore")
).hexdigest()[:16]


def get_llm_key(x_llm_key: Optional[str] = Header(None)) -> Optional[str]:
    """Optional per-request LLM key override provided by the caller via header 'X-LLM-Key'."""
    return x_llm_key


def _ensure_memory_file() -> None:
    """Fallback to JSON file if PostgreSQL not available"""
    if not MEMORY_PATH.exists():
        MEMORY_PATH.write_text(json.dumps({"items": []}), encoding="utf-8")


async def _load_memory_async() -> Dict[str, Any]:
    """Load memory - prefer PostgreSQL, fallback to JSON file"""
    # Try PostgreSQL first
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            doc_count = await mem_system.get_document_count()
            return {
                "type": "postgresql_pgvector",
                "database_url": os.getenv("DATABASE_URL", "not_set"),
                "embedding_dimension": EMBED_DIM,
                "available": True,
                "document_count": doc_count
            }
    except Exception as e:
        print(f"PostgreSQL memory system failed: {e}")
    
    # Fallback to JSON file
    _ensure_memory_file()
    try:
        return {
            "type": "json_file",
            "path": str(MEMORY_PATH),
            "items": json.loads(MEMORY_PATH.read_text(encoding="utf-8")).get("items", [])
        }
    except Exception:
        return {"type": "json_file", "path": str(MEMORY_PATH), "items": []}

def _load_memory() -> Dict[str, Any]:
    """Synchronous wrapper for memory loading"""
    try:
        # Try to get async result in sync context
        import asyncio
        try:
            # Check if we're already in an event loop
            loop = asyncio.get_running_loop()
            # We're in an async context, can't use run_until_complete
            return {
                "type": "postgresql_pgvector",
                "database_url": os.getenv("DATABASE_URL", "not_set"),
                "embedding_dimension": EMBED_DIM,
                "available": True,
                "document_count": 0  # Will be updated by async calls
            }
        except RuntimeError:
            # No event loop running, safe to create one
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                result = loop.run_until_complete(_load_memory_async())
                return result
            finally:
                loop.close()
    except Exception as e:
        print(f"Memory loading failed: {e}")
    
    # Fallback to JSON file
    _ensure_memory_file()
    try:
        return {
            "type": "json_file",
            "path": str(MEMORY_PATH),
            "items": json.loads(MEMORY_PATH.read_text(encoding="utf-8")).get("items", [])
        }
    except Exception:
        return {"type": "json_file", "path": str(MEMORY_PATH), "items": []}


def _save_memory(data: Dict[str, Any]) -> None:
    """Save memory - prefer PostgreSQL, fallback to JSON file"""
    # Try PostgreSQL first
    try:
        mem_system = _get_memory_system_sync()
        if mem_system:
            try:
                # PostgreSQL memory is handled by the MemorySystem class
                return
            except Exception:
                pass
    except Exception:
        pass
    
    # Fallback to JSON file
    MEMORY_PATH.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")


def _get_openai_client(override_key: Optional[str] = None) -> Optional[OpenAI]:
    api_key = override_key or os.getenv("OPENAI_API_KEY")
    if not api_key or OpenAI is None:
        return None
    return OpenAI(api_key=api_key)


def _ollama_post(path: str, payload: Dict[str, Any], timeout: int = 120) -> Dict[str, Any]:
    url = f"{OLLAMA_BASE}{path}"
    r = requests.post(url, json=payload, timeout=timeout)
    r.raise_for_status()
    return r.json()

def _ollama_get(path: str, timeout: int = 10) -> Dict[str, Any]:
    url = f"{OLLAMA_BASE}{path}"
    r = requests.get(url, timeout=timeout)
    r.raise_for_status()
    return r.json()


def _embed_texts_openai(client: OpenAI, texts: List[str]) -> List[List[float]]:
    res = client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [d.embedding for d in res.data]


def _embed_texts_ollama(texts: List[str]) -> List[List[float]]:
    vectors: List[List[float]] = []
    for t in texts:
        data = _ollama_post(
            "/api/embeddings",
            {"model": OLLAMA_EMBED_MODEL, "prompt": t},
            timeout=300,
        )
        # Ollama returns {"embedding": [...]}
        vec = data.get("embedding") or data.get("embeddings")
        if isinstance(vec, list):
            vectors.append(vec)
        else:
            vectors.append([0.0] * EMBED_DIM)
    return vectors


def _cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    if a.size == 0 or b.size == 0:
        return 0.0
    denom = (np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def _embed_texts(client: Optional[OpenAI], texts: List[str]) -> List[List[float]]:
    if LLM_PROVIDER == "ollama":
        return _embed_texts_ollama(texts)
    if client is None:
        return []
    return _embed_texts_openai(client, texts)


def _retrieve(client: Optional[OpenAI], query: str, k: int = 5) -> List[Dict[str, Any]]:
    mem = _load_memory()
    items: List[Dict[str, Any]] = mem.get("items", [])
    if not items:
        return []
    embs = _embed_texts(client, [query])
    if not embs:
        return []
    q_emb = np.array(embs[0], dtype=np.float32)
    scored: List[Tuple[float, Dict[str, Any]]] = []
    for item in items:
        emb = np.array(item.get("embedding", []), dtype=np.float32)
        scored.append((_cosine_sim(q_emb, emb), item))
    scored.sort(key=lambda t: t[0], reverse=True)
    return [it for _, it in scored[:k]]


# LLM raw generate
class GenerateBody(BaseModel):
    prompt: str


@app.post("/tools/llm/generate")
def llm_generate(body: GenerateBody, _: None = Depends(require_api_key), llm_key: Optional[str] = Depends(get_llm_key)) -> Dict[str, Any]:
    if LLM_PROVIDER == "ollama":
        try:
            data = _ollama_post(
                "/api/generate",
                {"model": OLLAMA_CHAT_MODEL, "prompt": body.prompt, "stream": False},
                timeout=600,
            )
            return {"response": (data.get("response") or "").strip()}
        except Exception as e:
            return {"response": f"[ollama error] {e}"}
    client = _get_openai_client(llm_key)
    if client is None:
        return {"response": f"[mock] You said: {body.prompt}"}
    msg = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": body.prompt},
    ]
    res = client.chat.completions.create(model=CHAT_MODEL, messages=msg)
    text = (res.choices[0].message.content or "").strip()
    return {"response": text}


# Chat and session
class ChatMessage(BaseModel):
    session_id: str
    message: str
    top_k: int = 5


@app.post("/chat/message")
async def chat_message(body: ChatMessage, _: None = Depends(require_api_key), llm_key: Optional[str] = Depends(get_llm_key)) -> Dict[str, Any]:
    # Tool use: allow explicit or inferred tool invocation
    tool_reply = await _handle_chat_tools(body.message)
    if tool_reply is not None:
        # Return tool result directly, still include used_docs=0 for UX consistency
        return {"answer": tool_reply["text"], "tool": tool_reply["tool"], "data": tool_reply["data"], "used_docs": 0}

    # Try to retrieve context using PostgreSQL memory system first
    contexts = []
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            # Use PostgreSQL memory system
            if LLM_PROVIDER == "ollama":
                try:
                    embed_data = _ollama_post(
                        "/api/embeddings",
                        {"model": OLLAMA_EMBED_MODEL, "prompt": body.message},
                        timeout=300
                    )
                    query_embedding = embed_data.get("embedding") or embed_data.get("embeddings", [])
                except Exception as e:
                    return {"answer": f"Embedding generation failed: {e}"}
            else:
                # OpenAI embedding
                client = _get_openai_client(llm_key)
                if client:
                    try:
                        embed_res = client.embeddings.create(
                            model=EMBED_MODEL,
                            input=body.message
                        )
                        query_embedding = embed_res.data[0].embedding
                    except Exception as e:
                        return {"answer": f"OpenAI embedding failed: {e}"}
                else:
                    query_embedding = []
            
            if query_embedding:
                # Query PostgreSQL
                try:
                    result = await mem_system.query_similar_with_embedding(body.message, query_embedding, k=body.top_k)
                    if result.get("success"):
                        contexts = result.get("snippets", [])
                except Exception as e:
                    # Fall back to old method
                    pass
    except Exception:
        # Fall back to old method
        pass
    
    # Fallback to old retrieval method if PostgreSQL failed
    if not contexts:
        client = _get_openai_client(llm_key) if LLM_PROVIDER != "ollama" else None
        contexts = _retrieve(client, body.message, k=max(1, min(10, body.top_k)))
    
    # Format context documents
    if contexts and hasattr(contexts[0], 'get'):
        # PostgreSQL results
        docs = "\n\n".join([f"[Doc {i+1}]\n" + (c.get("preview", "") or "") for i, c in enumerate(contexts)])
    else:
        # Old format results
        docs = "\n\n".join([f"[Doc {i+1}]\n" + (c.get("text", "") or "") for i, c in enumerate(contexts)])

    system = (
        "You are the agent brain with access to retrieved context from prior sessions. "
        "Use the provided documents if relevant. If not enough information is present, say so briefly."
    )
    user_msg = (
        f"User question:\n{body.message}\n\n"
        f"Retrieved context:\n{docs if docs else '(no context)'}"
    )

    if LLM_PROVIDER == "ollama":
        try:
            data = _ollama_post(
                "/api/generate",
                {
                    "model": OLLAMA_CHAT_MODEL,
                    "prompt": system + "\n\n" + user_msg,
                    "stream": False,
                    "options": {"temperature": 0.3},
                },
                timeout=600,
            )
            answer = (data.get("response") or "").strip()
            
            # Store chat history
            try:
                if mem_system:
                    await mem_system.store_chat_history(
                        session_id=body.session_id,
                        user_message=body.message,
                        assistant_response=answer,
                        context_docs=contexts,
                        used_docs=len(contexts),
                        metadata={"model": OLLAMA_CHAT_MODEL, "provider": "ollama"},
                        streamed=False
                    )
            except Exception as e:
                # Log error but don't fail the chat
                print(f"Failed to store chat history: {e}")
            
            return {"answer": answer, "used_docs": len(contexts)}
        except Exception as e:
            return {"answer": f"[ollama error] {e}"}

    # OpenAI path
    client = _get_openai_client(llm_key)
    if client is None:
        answer = f"[mock] LLM answer to: {body.message}"
        
        # Store chat history for mock responses too
        try:
            if mem_system:
                await mem_system.store_chat_history(
                    session_id=body.session_id,
                    user_message=body.message,
                    assistant_response=answer,
                    context_docs=contexts,
                    used_docs=len(contexts),
                    metadata={"model": "mock", "provider": "mock"},
                    streamed=False
                )
        except Exception as e:
            # Log error but don't fail the chat
            print(f"Failed to store chat history: {e}")
        
        return {"answer": answer, "used_docs": len(contexts)}
    
    res = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user_msg}],
        temperature=0.3,
    )
    answer = (res.choices[0].message.content or "").strip()
    
    # Store chat history
    try:
        if mem_system:
            await mem_system.store_chat_history(
                session_id=body.session_id,
                user_message=body.message,
                assistant_response=answer,
                context_docs=contexts,
                used_docs=len(contexts),
                metadata={"model": CHAT_MODEL, "provider": "openai"},
                streamed=False
            )
    except Exception as e:
        # Log error but don't fail the chat
        print(f"Failed to store chat history: {e}")
    
    return {"answer": answer, "used_docs": len(contexts)}


@app.post("/chat/stream")
async def chat_stream(body: ChatMessage, _: None = Depends(require_api_key), llm_key: Optional[str] = Depends(get_llm_key)):
    """Stream chat answer with periodic keepalives so tunnels/proxies don't time out.
    Sends small whitespace heartbeats every ~8 seconds and yields partial tokens
    as available. Client should treat it as text/event-stream or chunked text.
    """
    # Try to retrieve context using PostgreSQL memory system first
    contexts = []
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            # Use PostgreSQL memory system
            if LLM_PROVIDER == "ollama":
                try:
                    embed_data = _ollama_post(
                        "/api/embeddings",
                        {"model": OLLAMA_EMBED_MODEL, "prompt": body.message},
                        timeout=300
                    )
                    query_embedding = embed_data.get("embedding") or embed_data.get("embeddings", [])
                except Exception as e:
                    return {"answer": f"Embedding generation failed: {e}"}
            else:
                # OpenAI embedding
                client = _get_openai_client(llm_key)
                if client:
                    try:
                        embed_res = client.embeddings.create(
                            model=EMBED_MODEL,
                            input=body.message
                        )
                        query_embedding = embed_res.data[0].embedding
                    except Exception as e:
                        return {"answer": f"OpenAI embedding failed: {e}"}
                else:
                    query_embedding = []
            
            if query_embedding:
                # Query PostgreSQL
                try:
                    result = await mem_system.query_similar_with_embedding(body.message, query_embedding, k=body.top_k)
                    if result.get("success"):
                        contexts = result.get("snippets", [])
                except Exception as e:
                    # Fall back to old method
                    pass
    except Exception:
        # Fall back to old method
        pass
    
    # Fallback to old retrieval method if PostgreSQL failed
    if not contexts:
        client = _get_openai_client(llm_key) if LLM_PROVIDER != "ollama" else None
        contexts = _retrieve(client, body.message, k=max(1, min(10, body.top_k)))
    
    # Format context documents
    if contexts and hasattr(contexts[0], 'get'):
        # PostgreSQL results
        docs = "\n\n".join([f"[Doc {i+1}]\n" + (c.get("preview", "") or "") for i, c in enumerate(contexts)])
    else:
        # Old format results
        docs = "\n\n".join([f"[Doc {i+1}]\n" + (c.get("text", "") or "") for i, c in enumerate(contexts)])
    
    system = (
        "You are the agent brain with access to retrieved context from prior sessions. "
        "Use the provided documents if relevant. If not enough information is present, say so briefly."
    )
    user_msg = (
        f"User question:\n{body.message}\n\n" f"Retrieved context:\n{docs if docs else '(no context)'}"
    )

    def _stream_ollama():
        import datetime
        last = time.time()
        full_response = ""
        try:
            # Use generate streaming
            url = f"{OLLAMA_BASE}/api/generate"
            with requests.post(
                url,
                json={
                    "model": OLLAMA_CHAT_MODEL,
                    "prompt": system + "\n\n" + user_msg,
                    "stream": True,
                    "options": {"temperature": 0.3},
                },
                stream=True,
                timeout=600,
            ) as r:
                r.raise_for_status()
                for line in r.iter_lines(decode_unicode=True):
                    now = time.time()
                    if not line:
                        # emit keepalive every 8s
                        if now - last > 8:
                            last = now
                            yield " \n"
                        continue
                    # Ollama streams JSON per line {response: '...', done: bool}
                    try:
                        obj = json.loads(line)
                        chunk = (obj.get("response") or "")
                        if chunk:
                            full_response += chunk
                            yield chunk
                            last = now
                        if obj.get("done"):
                            break
                    except Exception:
                        yield " \n"
                
                # Store chat history after streaming is complete
                try:
                    if mem_system and full_response.strip():
                        # Chat history will be stored by the wrapper function
                        pass
                except Exception as e:
                    print(f"Failed to store streamed chat history: {e}")
                    
        except Exception as e:
            yield f"[stream error] {e}"

    def _stream_openai():
        last = time.time()
        full_response = ""
        if client is None:
            # mock stream with small chunks
            txt = f"[mock] LLM answer to: {body.message}"
            for ch in txt:
                full_response += ch
                yield ch
                time.sleep(0.01)
            
            # Store chat history for mock responses
            try:
                if mem_system:
                    # Chat history will be stored by the wrapper function
                    pass
            except Exception as e:
                print(f"Failed to store mock streamed chat history: {e}")
            return
            
        try:
            stream = client.chat.completions.create(
                model=CHAT_MODEL,
                messages=[{"role": "system", "content": system}, {"role": "user", "content": user_msg}],
                temperature=0.3,
                stream=True,
            )
            for ev in stream:
                now = time.time()
                if hasattr(ev, "choices") and ev.choices:
                    delta = ev.choices[0].delta.content or ""
                    if delta:
                        full_response += delta
                        yield delta
                        last = now
                if now - last > 8:
                    yield " \n"
                    last = now
            
            # Store chat history after streaming is complete
            try:
                if mem_system and full_response.strip():
                    # Chat history will be stored by the wrapper function
                    pass
            except Exception as e:
                print(f"Failed to store streamed chat history: {e}")
                
        except Exception as e:
            yield f"[stream error] {e}"

    generator = _stream_ollama if LLM_PROVIDER == "ollama" else _stream_openai
    
    # Create a wrapper generator that captures the full response and stores it
    async def _stream_with_history():
        full_response = ""
        try:
            for chunk in generator():
                full_response += chunk
                yield chunk
        finally:
            # Store chat history after streaming is complete
            try:
                if mem_system and full_response.strip():
                    await mem_system.store_chat_history(
                        session_id=body.session_id,
                        user_message=body.message,
                        assistant_response=full_response.strip(),
                        context_docs=contexts,
                        used_docs=len(contexts),
                        metadata={"model": OLLAMA_CHAT_MODEL if LLM_PROVIDER == "ollama" else CHAT_MODEL, "provider": LLM_PROVIDER},
                        streamed=True
                    )
            except Exception as e:
                print(f"Failed to store streamed chat history: {e}")
    
    return StreamingResponse(_stream_with_history(), media_type="text/plain; charset=utf-8")

class RunSessionBody(BaseModel):
    duration_min: int = 2
    sleep: float = 0.2
    use_llm: bool = True
    preload_llm: bool = True
    max_tests_per_loop: int = 1


@app.post("/admin/run_session")
def run_session(_: RunSessionBody, __: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"returncode": 0, "stdout": "[mock] session completed", "stderr": ""}


# ---------------------------
# Readiness + Warmup
# ---------------------------

class WarmupResult(BaseModel):
    provider: str
    embed_ms: Optional[float] = None
    generate_ms: Optional[float] = None
    ready: bool = False


@app.get("/status/ready")
def status_ready(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    ready = False
    details: Dict[str, Any] = {"provider": LLM_PROVIDER}
    if LLM_PROVIDER == "ollama":
        try:
            _ = _ollama_get("/api/tags", timeout=5)
            details["ollama"] = "up"
            ready = True
        except Exception as e:
            details["ollama_error"] = str(e)
            ready = False
    else:
        details["has_openai_key"] = bool(os.getenv("OPENAI_API_KEY"))
        ready = bool(os.getenv("OPENAI_API_KEY"))
    return {"ready": ready, "details": details}


@app.post("/admin/warmup", response_model=WarmupResult)
def admin_warmup(_: None = Depends(require_api_key), llm_key: Optional[str] = Depends(get_llm_key)) -> WarmupResult:
    embed_ms: Optional[float] = None
    generate_ms: Optional[float] = None
    if LLM_PROVIDER == "ollama":
        try:
            t0 = time.time()
            _ = _ollama_post("/api/embeddings", {"model": OLLAMA_EMBED_MODEL, "prompt": "warmup"}, timeout=180)
            embed_ms = (time.time() - t0) * 1000.0
        except Exception:
            embed_ms = None
        try:
            t0 = time.time()
            _ = _ollama_post("/api/generate", {"model": OLLAMA_CHAT_MODEL, "prompt": "hi", "stream": False}, timeout=300)
            generate_ms = (time.time() - t0) * 1000.0
        except Exception:
            generate_ms = None
        ready = (embed_ms is not None) and (generate_ms is not None)
        return WarmupResult(provider=LLM_PROVIDER, embed_ms=embed_ms, generate_ms=generate_ms, ready=ready)

    # OpenAI path
    client = _get_openai_client(llm_key)
    if client is None:
        return WarmupResult(provider=LLM_PROVIDER, embed_ms=None, generate_ms=None, ready=False)
    try:
        t0 = time.time()
        _ = _embed_texts_openai(client, ["warmup text"])
        embed_ms = (time.time() - t0) * 1000.0
    except Exception:
        embed_ms = None
    try:
        t0 = time.time()
        _ = client.chat.completions.create(model=CHAT_MODEL, messages=[{"role": "user", "content": "hi"}])
        generate_ms = (time.time() - t0) * 1000.0
    except Exception:
        generate_ms = None
    ready = (embed_ms is not None) and (generate_ms is not None)
    return WarmupResult(provider=LLM_PROVIDER, embed_ms=embed_ms, generate_ms=generate_ms, ready=ready)

# Memory store/query (JSON + OpenAI embeddings)
class StoreBody(BaseModel):
    user_input: str
    response: Optional[str] = None
    response_b64: Optional[str] = None


@app.post("/memory/store")
async def memory_store(body: StoreBody, _: None = Depends(require_api_key), llm_key: Optional[str] = Depends(get_llm_key)) -> Dict[str, Any]:
    """Store user input and response in memory with embeddings"""
    try:
        # Try PostgreSQL memory system first
        mem_system = await _get_simple_memory_system()
        if mem_system:
            try:
                # Use PostgreSQL memory system
                from src.layers.memory_system import Interaction
                from datetime import datetime
                
                # Create interaction object
                interaction = Interaction(
                    user_input=body.user_input,
                    response=body.response,
                    tool_calls=[],
                    metadata={},
                    timestamp=datetime.now()
                )
                
                # Generate embedding
                if LLM_PROVIDER == "ollama":
                    try:
                        embed_data = _ollama_post(
                            "/api/embeddings",
                            {"model": OLLAMA_EMBED_MODEL, "prompt": body.user_input + "\n" + body.response},
                            timeout=300
                        )
                        embedding = embed_data.get("embedding") or embed_data.get("embeddings", [])
                    except Exception as e:
                        return {"success": False, "error": f"Embedding generation failed: {e}"}
                else:
                    # OpenAI embedding
                    client = _get_openai_client(llm_key)
                    if client:
                        try:
                            embed_res = client.embeddings.create(
                                model=EMBED_MODEL,
                                input=body.user_input + "\n" + body.response
                            )
                            embedding = embed_res.data[0].embedding
                        except Exception as e:
                            return {"success": False, "error": f"OpenAI embedding failed: {e}"}
                    else:
                        return {"success": False, "error": "No embedding provider available"}
                
                # Store in PostgreSQL
                try:
                    # Preserve original filename in content if provided via metadata
                    result = await mem_system.store_interaction(interaction, embedding)
                    if result.get("success"):
                        return {
                            "success": True,
                            "message": "Stored in PostgreSQL vector database",
                            "record_id": result.get("record_id"),
                            "type": "postgresql_pgvector",
                        }
                    else:
                        return {
                            "success": False,
                            "error": f"PostgreSQL storage failed: {result.get('error')}",
                        }
                except Exception as e:
                    return {"success": False, "error": f"PostgreSQL operation failed: {e}"}
                
            except Exception as e:
                return {"success": False, "error": f"PostgreSQL memory system error: {e}"}
        
        # Fallback to JSON file storage
        mem = _load_memory()
        items = mem.get("items", [])
        
        # Generate embedding for similarity search
        if LLM_PROVIDER == "ollama":
            try:
                embed_data = _ollama_post(
                    "/api/embeddings",
                    {"model": OLLAMA_EMBED_MODEL, "prompt": body.user_input + "\n" + body.response},
                    timeout=300
                )
                embedding = embed_data.get("embedding") or embed_data.get("embeddings", [])
            except Exception as e:
                return {"success": False, "error": f"Embedding generation failed: {e}"}
        else:
            # OpenAI embedding
            client = _get_openai_client(llm_key)
            if client:
                try:
                    embed_res = client.embeddings.create(
                        model=EMBED_MODEL,
                        input=body.user_input + "\n" + body.response
                    )
                    embedding = embed_res.data[0].embedding
                except Exception as e:
                    return {"success": False, "error": f"OpenAI embedding failed: {e}"}
            else:
                embedding = []
        
        # Store in JSON file
        item = {
            "user_input": body.user_input,
            "response": body.response,
            "embedding": embedding,
            "timestamp": time.time(),
            "metadata": {}
        }
        items.append(item)
        _save_memory({"items": items})
        
        return {
            "success": True,
            "message": "Stored in JSON file (PostgreSQL not available)",
            "type": "json_file",
            "total_items": len(items)
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/memory/query")
async def memory_query(query: str, k: int = 5, _: None = Depends(require_api_key), llm_key: Optional[str] = Depends(get_llm_key)) -> Dict[str, Any]:
    """Query memory for similar content using simple PostgreSQL search"""
    try:
        # Try simple PostgreSQL memory system first
        mem_system = await _get_simple_memory_system()
        if mem_system:
            try:
                results = await mem_system.query_similar(query, k)
                return {
                    "success": True,
                    "type": "postgresql_pgvector",
                    "query": query,
                    "total_found": len(results),
                    "snippets": results
                }
            except Exception as e:
                print(f"Simple PostgreSQL query failed: {e}")
        
        # Fallback to JSON file storage
        mem = _load_memory()
        items = mem.get("items", [])
        if not items:
            return {"success": True, "type": "json_file", "query": query, "snippets": [], "total_found": 0}
        
        # Simple text search in JSON items
        results = []
        for item in items:
            if query.lower() in item.get("user_input", "").lower() or query.lower() in item.get("response", "").lower():
                results.append(item)
                if len(results) >= k:
                    break
        
        return {
            "success": True,
            "type": "json_file",
            "query": query,
            "total_found": len(results),
            "snippets": results
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/memory/documents")
async def get_available_documents(limit: int = Query(10, description="Number of documents to return", ge=1, le=50), _: None = Depends(require_api_key)) -> Dict[str, Any]:
    """Get available documents for preview"""
    try:
        # Try simple PostgreSQL memory system first
        mem_system = await _get_simple_memory_system()
        if mem_system:
            try:
                documents = await mem_system.get_sample_documents(limit)
                doc_count = await mem_system.get_document_count()
                return {
                    "success": True,
                    "type": "postgresql_pgvector",
                    "total_documents": doc_count,
                    "documents": documents
                }
            except Exception as e:
                print(f"Failed to get documents from PostgreSQL: {e}")
        
        # Fallback to JSON file
        mem = _load_memory()
        items = mem.get("items", [])
        return {
            "success": True,
            "type": "json_file",
            "total_documents": len(items),
            "documents": items[:limit]
        }
        
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.get("/memory/document/{doc_id}")
async def get_document(doc_id: int, _: None = Depends(require_api_key)) -> Dict[str, Any]:
    """Fetch a single document by id with full content."""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            doc = await mem_system.get_document_by_id(doc_id)
            if doc:
                return {"success": True, "type": "postgresql_pgvector", "document": doc}
            return {"success": False, "error": "not_found"}
        # Fallback JSON: derive from list index
        mem = _load_memory()
        items = mem.get("items", [])
        idx = max(0, min(len(items) - 1, doc_id - 1))
        if 0 <= idx < len(items):
            it = items[idx]
            return {"success": True, "type": "json_file", "document": {"id": str(doc_id), "source": "json_file", "content": it.get("response", ""), "created_at": None}}
        return {"success": False, "error": "not_found"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# Content preview (mock)
@app.get("/preview/file")
def preview_file(file_path: str, preview_type: Optional[str] = None, _: None = Depends(require_api_key)) -> Dict[str, Any]:
    content = f"[mock preview] {file_path} ({preview_type or 'auto'})"
    return {"file_type": preview_type or "text", "language": "markdown", "content_length": len(content), "content": content, "metadata": {"source_type": "mock"}, "preview_html": f"<pre>{content}</pre>"}


@app.get("/preview/analyze")
def preview_analyze(file_path: str, _: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {
        "preview_type": "markdown",
        "language": "markdown",
        "supported": True,
        "source_type": "mock",
        "file_extension": ".md",
        "exists": True,
        "capabilities": {"syntax_highlighting": True, "markdown_rendering": True},
    }


@app.get("/preview/supported-types")
def preview_supported_types(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return {"supported_types": {"code": {"extensions": [".py", ".js"], "mime_types": ["text/plain"]}, "markdown": {"extensions": [".md"], "mime_types": ["text/markdown"]}}}


class BatchBody(BaseModel):
    file_paths: List[str]


# ---------------------------
# Admin: backfill source from content
# ---------------------------
@app.post("/admin/backfill_sources")
async def admin_backfill_sources(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    """Backfill `source` with filename if it is embedded in the content.

    Helps older records where original filename wasn't stored in `source`.
    Looks for a prefix like `upload:<name>` or `[filename: <name>]`.
    """
    try:
        mem_system = await _get_simple_memory_system()
        if not mem_system:
            return {"success": False, "error": "Memory system not available"}

        if not mem_system._pool:
            await mem_system.initialize()

        updated = 0
        async with mem_system._pool.acquire() as conn:  # type: ignore[attr-defined]
            rows = await conn.fetch(
                """
                SELECT id, source, content FROM vector_store
                ORDER BY id DESC LIMIT 2000
                """
            )
            for r in rows:
                src = r["source"] or ""
                if src and src != "unknown":
                    continue
                content = r["content"] or ""
                m = re.search(r"\[filename:\s*([^\]\n]+)\]", content)
                if not m:
                    m = re.search(r"upload:([^\s]+)", content)
                if m:
                    filename = m.group(1).strip()
                    await conn.execute("UPDATE vector_store SET source=$1 WHERE id=$2", filename, r["id"])
                    updated += 1
        return {"success": True, "updated": updated}
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.post("/preview/batch")
def preview_batch(body: BatchBody, _: None = Depends(require_api_key)) -> Dict[str, Any]:
    results = []
    for p in body.file_paths:
        content = f"[mock batch preview] {p}"
        results.append({"file_path": p, "success": True, "file_type": "text", "language": "markdown", "content_length": len(content), "preview_html": f"<pre>{content}</pre>"})
    return {"successful_previews": len(results), "total_files": len(body.file_paths), "results": results}


# ---------------------------
# Component status snapshot
# ---------------------------

@app.get("/status/components")
def status_components(_: None = Depends(require_api_key)) -> Dict[str, Any]:
    return _components_snapshot()


def _components_snapshot() -> Dict[str, Any]:
    components: Dict[str, Any] = {}
    # API basics
    components["api"] = {
        "service": "capstone-demo-api",
        "llm_provider": LLM_PROVIDER,
        "hostname": HOSTNAME,
        "started_at": STARTED_AT,
        "image_tag": IMAGE_TAG,
        "commit_sha": COMMIT_SHA,
        "build_time": BUILD_TIME,
    }
    # Endpoints availability
    components["endpoints"] = {
        "status_ready": True,
        "admin_warmup": True,
        "chat_stream": True,
        "ingest_files": True,
    }
    # Memory system status
    try:
        mem_info = _load_memory()
        if mem_info.get("type") == "postgresql_pgvector":
            # PostgreSQL memory system
            components["memory"] = {
                "type": "postgresql_pgvector",
                "database_url": mem_info.get("database_url", "not_set"),
                "embedding_dimension": mem_info.get("embedding_dimension", EMBED_DIM),
                "available": mem_info.get("available", False),
                "status": "active" if mem_info.get("available") else "error"
            }
        else:
            # JSON file fallback
            components["memory"] = {
                "type": "json_file",
                "path": str(MEMORY_PATH),
                "items": mem_info.get("items", []),
                "status": "fallback"
            }
    except Exception as e:
        components["memory"] = {"error": str(e), "status": "error"}
    
    # LLM provider quick signal
    if LLM_PROVIDER == "ollama":
        try:
            _ = _ollama_get("/api/tags", timeout=5)
            components["llm"] = {"ollama": "up", "base": OLLAMA_BASE, "chat_model": OLLAMA_CHAT_MODEL, "embed_model": OLLAMA_EMBED_MODEL}
        except Exception as e:
            components["llm"] = {"ollama": "down", "error": str(e)}
    else:
        components["llm"] = {"openai_key": bool(os.getenv("OPENAI_API_KEY")), "chat_model": CHAT_MODEL, "embed_model": EMBED_MODEL}
    return components


@app.get("/")
def root() -> Dict[str, Any]:
    return {"service": "capstone-demo-api", "status": "ok"}


# ---------------------------
# Ingestion utilities
# ---------------------------

def _extract_text_from_bytes(content: bytes, filename: str) -> str:
    name = (filename or "").lower()
    try:
        if name.endswith(('.md', '.txt', '.json', '.csv', '.log', '.yaml', '.yml', '.py', '.js', '.ts', '.html', '.htm')):
            return content.decode('utf-8', errors='ignore')
        if name.endswith('.pdf'):
            from io import BytesIO
            from pdfminer.high_level import extract_text
            return extract_text(BytesIO(content)) or ""
        if name.endswith('.docx'):
            from io import BytesIO
            import docx
            doc = docx.Document(BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])
        if name.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
            try:
                import pytesseract  # optional if available
                from PIL import Image
                from io import BytesIO
                img = Image.open(BytesIO(content))
                return pytesseract.image_to_string(img) or ""
            except Exception:
                return ""
    except Exception:
        return ""
    # Fallback
    return content.decode('utf-8', errors='ignore')


class IngestResponse(BaseModel):
    stored: int
    errors: int
    details: List[Dict[str, Any]]


@app.post("/ingest/files", response_model=IngestResponse)
async def ingest_files(
    files: List[UploadFile] = File(...),
    source: Optional[str] = Form(None),
    _: None = Depends(require_api_key),
    llm_key: Optional[str] = Depends(get_llm_key),
):
    stored = 0
    errs = 0
    details: List[Dict[str, Any]] = []
    
    for uf in files:
        try:
            data = await uf.read()
            text = _extract_text_from_bytes(data, uf.filename or "file")
            if not text.strip():
                details.append({"file": uf.filename, "ok": False, "reason": "empty"})
                errs += 1
                continue
            
            # Try to store in PostgreSQL memory system first
            try:
                mem_system = await _get_simple_memory_system()
                if mem_system:
                    # Use PostgreSQL memory system
                    from src.layers.memory_system import Interaction
                    from datetime import datetime
                    
                    # Create interaction object
                    original_name = (uf.filename or "file")
                    interaction = Interaction(
                        user_input=f"upload:{original_name} {source or ''}",
                        response=f"[filename: {original_name}]\n{text}",
                        tool_calls=[],
                        metadata={"source": "file_upload", "filename": original_name},
                        timestamp=datetime.now()
                    )
                    
                    # Generate embedding
                    if LLM_PROVIDER == "ollama":
                        try:
                            embed_data = _ollama_post(
                                "/api/embeddings",
                                {"model": OLLAMA_EMBED_MODEL, "prompt": text},
                                timeout=300
                            )
                            embedding = embed_data.get("embedding") or embed_data.get("embeddings", [])
                        except Exception as e:
                            details.append({"file": uf.filename, "ok": False, "error": f"Embedding failed: {e}"})
                            errs += 1
                            continue
                    else:
                        # OpenAI embedding
                        client = _get_openai_client(llm_key)
                        if client:
                            try:
                                embed_res = client.embeddings.create(
                                    model=EMBED_MODEL,
                                    input=text
                                )
                                embedding = embed_res.data[0].embedding
                            except Exception as e:
                                details.append({"file": uf.filename, "ok": False, "error": f"OpenAI embedding failed: {e}"})
                                errs += 1
                                continue
                        else:
                            details.append({"file": uf.filename, "ok": False, "error": "No embedding provider available"})
                            errs += 1
                            continue
                    
                    # Store in PostgreSQL
                    try:
                        result = await mem_system.store_interaction(interaction, embedding)
                        if result.get("success"):
                            details.append({"file": uf.filename, "ok": True, "chars": len(text), "type": "postgresql_pgvector"})
                            stored += 1
                            continue
                        else:
                            details.append({"file": uf.filename, "ok": False, "error": f"PostgreSQL storage failed: {result.get('error')}"})
                            errs += 1
                            continue
                    except Exception as e:
                        details.append({"file": uf.filename, "ok": False, "error": f"PostgreSQL operation failed: {e}"})
                        errs += 1
                        continue
                        
            except Exception as e:
                # Fall back to old method
                pass
            
            # Fallback to old JSON file storage
            body = StoreBody(user_input=f"upload:{uf.filename} {source or ''}", response=text)
            await memory_store(body, None, llm_key)  # reuse in-process
            details.append({"file": uf.filename, "ok": True, "chars": len(text), "type": "json_file"})
            stored += 1
            
        except Exception as e:
            errs += 1
            details.append({"file": uf.filename, "ok": False, "error": str(e)})
    
    return IngestResponse(stored=stored, errors=errs, details=details)


# Chat History Endpoints
@app.get("/chat/history")
async def chat_history(session_id: str = Query(None, description="Session ID to filter by"), 
                      limit: int = Query(50, description="Maximum number of records to return", ge=1, le=100),
                      offset: int = Query(0, description="Number of records to skip", ge=0),
                      include_context: bool = Query(False, description="Include context documents"),
                      _: None = Depends(require_api_key)) -> Dict[str, Any]:
    """Get chat history from the database"""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            result = await mem_system.get_chat_history(
                session_id=session_id,
                limit=limit,
                offset=offset,
                include_context=include_context
            )
            return result
        else:
            return {
                "success": False,
                "error": "Memory system not available"
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/chat/history/{session_id}")
async def chat_history_by_session(session_id: str,
                                limit: int = Query(50, description="Maximum number of records to return", ge=1, le=100),
                                offset: int = Query(0, description="Number of records to skip", ge=0),
                                include_context: bool = Query(False, description="Include context documents"),
                                _: None = Depends(require_api_key)) -> Dict[str, Any]:
    """Get chat history for a specific session"""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            result = await mem_system.get_chat_history(
                session_id=session_id,
                limit=limit,
                offset=offset,
                include_context=include_context
            )
            return result
        else:
            return {
                "success": False,
                "error": "Memory system not available"
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/chat/session/{session_id}/summary")
async def chat_session_summary(session_id: str,
                             _: None = Depends(require_api_key)) -> Dict[str, Any]:
    """Get a summary of a chat session"""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            result = await mem_system.get_session_summary(session_id)
            return result
        else:
            return {
                "success": False,
                "error": "Memory system not available"
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/chat/sessions")
async def list_chat_sessions(limit: int = Query(20, description="Maximum number of sessions to return", ge=1, le=100),
                           offset: int = Query(0, description="Number of sessions to skip", ge=0),
                           _: None = Depends(require_api_key)) -> Dict[str, Any]:
    """List all available chat sessions"""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            # Get all chat history to extract unique sessions
            result = await mem_system.get_chat_history(limit=1000, offset=0)
            if not result.get("success"):
                return result
            
            # Extract unique sessions
            sessions = {}
            for chat in result.get("chats", []):
                session_id = chat["session_id"]
                if session_id not in sessions:
                    sessions[session_id] = {
                        "session_id": session_id,
                        "first_message": chat["created_at"],
                        "last_message": chat["created_at"],
                        "message_count": 0,
                        "total_docs_used": 0,
                        "streamed_messages": 0
                    }
                
                sessions[session_id]["message_count"] += 1
                sessions[session_id]["total_docs_used"] += chat["used_docs"]
                if chat["streamed"]:
                    sessions[session_id]["streamed_messages"] += 1
                
                # Update timestamps
                if chat["created_at"] < sessions[session_id]["first_message"]:
                    sessions[session_id]["first_message"] = chat["created_at"]
                if chat["created_at"] > sessions[session_id]["last_message"]:
                    sessions[session_id]["last_message"] = chat["created_at"]
            
            # Convert to list and sort by last message
            session_list = list(sessions.values())
            session_list.sort(key=lambda x: x["last_message"], reverse=True)
            
            # Apply pagination
            total_sessions = len(session_list)
            paginated_sessions = session_list[offset:offset + limit]
            
            return {
                "success": True,
                "sessions": paginated_sessions,
                "total_count": total_sessions,
                "limit": limit,
                "offset": offset,
                "has_more": (offset + limit) < total_sessions
            }
        else:
            return {
                "success": False,
                "error": "Memory system not available"
            }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

# Add enhanced ingestion imports
from local_enhanced_ingestion import LocalEnhancedIngestionEngine
from enhanced_storage import EnhancedStorageEngine

# Add enhanced ingestion endpoints
@app.post("/ingest/enhanced")
async def ingest_file_enhanced(
    file: UploadFile,
    extraction_options: Dict[str, Any] = None,
    _: None = Depends(require_api_key)
) -> Dict[str, Any]:
    """Enhanced file ingestion using local Ollama models"""
    
    try:
        # Get memory system
        mem_system = await _get_simple_memory_system()
        if not mem_system:
            return {
                "success": False,
                "error": "Memory system not available"
            }
        
        # Save uploaded file temporarily
        temp_file_path = f"temp_uploads/{file.filename}"
        os.makedirs("temp_uploads", exist_ok=True)
        
        with open(temp_file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Initialize enhanced ingestion engine
        async with LocalEnhancedIngestionEngine(
            ollama_base_url=OLLAMA_BASE,
            current_system=mem_system
        ) as engine:
            
            # Process file with enhanced ingestion
            ingestion_results = await engine.ingest_file_enhanced(
                temp_file_path,
                file_type=file.filename.split('.')[-1] if '.' in file.filename else None,
                extraction_options=extraction_options
            )
            
            # Store results using enhanced storage
            storage_engine = EnhancedStorageEngine(
                current_memory_system=mem_system,
                database_url=os.getenv("DATABASE_URL")
            )
            
            storage_results = await storage_engine.store_document_enhanced(
                ingestion_results
            )
            
            # Clean up temp file
            try:
                os.remove(temp_file_path)
            except:
                pass
            
            return {
                "success": True,
                "ingestion": ingestion_results,
                "storage": storage_results,
                "file_info": {
                    "filename": file.filename,
                    "size_bytes": len(content),
                    "content_type": file.content_type
                }
            }
    
    except Exception as e:
        logger.error(f"Enhanced ingestion failed: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }

@app.post("/ingest/enhanced/batch")
async def ingest_files_enhanced_batch(
    files: List[UploadFile],
    extraction_options: Dict[str, Any] = None,
    _: None = Depends(require_api_key)
) -> Dict[str, Any]:
    """Batch enhanced file ingestion"""
    
    try:
        results = []
        
        for file in files:
            # Process each file individually
            file_result = await ingest_file_enhanced(file, extraction_options, _)
            results.append({
                "filename": file.filename,
                "result": file_result
            })
        
        # Generate batch summary
        successful = sum(1 for r in results if r["result"]["success"])
        failed = len(results) - successful
        
        return {
            "success": True,
            "batch_summary": {
                "total_files": len(files),
                "successful": successful,
                "failed": failed
            },
            "results": results
        }
    
    except Exception as e:
        logger.error(f"Batch enhanced ingestion failed: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }

@app.get("/ingest/enhanced/status")
async def get_enhanced_ingestion_status(
    _: None = Depends(require_api_key)
) -> Dict[str, Any]:
    """Get status of enhanced ingestion system"""
    
    try:
        # Check Ollama availability
        ollama_health = await check_ollama_health()
        
        # Check current system availability
        current_system_health = await check_current_system_health()
        
        return {
            "success": True,
            "enhanced_ingestion": {
                "status": "available" if ollama_health["healthy"] else "unavailable",
                "ollama_models": ollama_health.get("models", []),
                "ollama_endpoint": OLLAMA_BASE
            },
            "current_system": {
                "status": "available" if current_system_health["healthy"] else "unavailable",
                "health": current_system_health
            },
            "fallback_available": True,  # Always available
            "hybrid_capabilities": {
                "structured_extraction": True,
                "entity_recognition": True,
                "document_analysis": True,
                "multi_format_support": True
            }
        }
    
    except Exception as e:
        logger.error(f"Failed to get enhanced ingestion status: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }

async def check_ollama_health() -> Dict[str, Any]:
    """Check Ollama service health"""
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(f"{OLLAMA_BASE}/api/tags", timeout=5) as response:
                if response.status == 200:
                    data = await response.json()
                    models = [model["name"] for model in data.get("models", [])]
                    return {
                        "healthy": True,
                        "models": models,
                        "endpoint": OLLAMA_BASE
                    }
                else:
                    return {
                        "healthy": False,
                        "error": f"HTTP {response.status}",
                        "endpoint": OLLAMA_BASE
                    }
    except Exception as e:
        return {
            "healthy": False,
            "error": str(e),
            "endpoint": OLLAMA_BASE
        }

async def check_current_system_health() -> Dict[str, Any]:
    """Check current memory system health"""
    try:
        mem_system = await _get_simple_memory_system()
        if mem_system:
            health = await mem_system.get_health_status()
            return {
                "healthy": health.get("status") == "healthy",
                "details": health
            }
        else:
            return {
                "healthy": False,
                "error": "Memory system not initialized"
            }
    except Exception as e:
        return {
            "healthy": False,
            "error": str(e)
        }



